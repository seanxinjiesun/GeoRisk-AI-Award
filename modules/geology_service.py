from __future__ import annotations

import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List

import pdfplumber
import pytesseract
from docx import Document
from PIL import Image

from modules.claude_client import call_claude


MINING_KEYWORDS = [
    "caf2",
    "sio2",
    "ore",
    "deposit",
    "grade",
    "矿",
    "矿体",
    "品位",
    "成矿",
    "钻孔",
    "萤石",
]


@dataclass
class FileAIResult:
    mode: str  # mining | general
    summary: str
    key_insights: List[str]
    bullet_points: List[str]
    risk_issues: str
    application_advice: str

    mineral_type: str
    grade_info: str
    deposit_type: str
    orebody_scale: str
    thickness_extension: str
    mineability: str

    geological_info: str
    orebody_analysis: str
    data_integrity: str
    risk_identification: str
    investment_advice: str

    result_interpretation: str
    logic_basis: str
    risk_hint: str
    risk_level: str
    radar_scores: Dict[str, float]
    highlight_keywords: List[str]
    raw_response: str


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_docx(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def _read_pdf_with_pdfplumber(path: Path) -> str:
    parts: List[str] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    return "\n".join(parts)


def _ocr_pdf_with_pytesseract(path: Path) -> str:
    try:
        import pypdfium2 as pdfium
    except Exception:
        return ""

    ocr_parts: List[str] = []
    try:
        pdf = pdfium.PdfDocument(str(path))
        for i in range(len(pdf)):
            page = pdf.get_page(i)
            bitmap = page.render(scale=2)
            pil_img = bitmap.to_pil()
            ocr_parts.append(pytesseract.image_to_string(pil_img, lang="eng+chi_sim"))
            page.close()
    except Exception:
        return ""

    return "\n".join(ocr_parts)


def _read_image_with_ocr(path: Path) -> str:
    try:
        img = Image.open(str(path))
        return pytesseract.image_to_string(img, lang="eng+chi_sim")
    except Exception:
        return ""


def extract_text(file_path: str) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return _read_txt(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix == ".pdf":
        text = _read_pdf_with_pdfplumber(path)
        if len(text.strip()) < 120:
            text = f"{text}\n{_ocr_pdf_with_pytesseract(path)}".strip()
        return text
    if suffix in {".png", ".jpg", ".jpeg", ".bmp"}:
        return _read_image_with_ocr(path)

    raise ValueError("仅支持 .txt/.docx/.pdf/.png/.jpg/.jpeg/.bmp 文件")


def _is_mining_related(report_text: str) -> bool:
    lower = report_text.lower()
    return any(k in lower for k in MINING_KEYWORDS)


def _build_prompt(mode: str, company_name: str, project_name: str, report_text: str) -> str:
    clipped = report_text[:12000]
    hard_rules = """
【强制执行指令】
1) 绝对忠于原文：必须仅基于文档内容提取和扩写，严禁捏造；如文档缺失某项信息，必须明确写“文件中未提及”。
2) 拒绝空泛总结：必须像资深矿业工程师和贸易风控专家，逐项深挖地质构造特征、成矿背景、核心金属品位具体数值、储量表格要点、选冶工艺难点、投资与贸易风险。
3) 结构化长文输出：必须在JSON各文本字段内使用Markdown标题/加粗/列表，信息尽量详尽，整体内容深度尽量接近4000字级别。
""".strip()

    if mode == "mining":
        return f"""
你是矿业投资与地质尽调专家。请输出严格JSON，不要输出JSON外文本。

{hard_rules}

公司名称：{company_name or '未提供'}
项目名称：{project_name or '未提供'}
文档内容：
{clipped}

输出JSON：
{{
  "summary": "整体结论摘要",
  "mineral_type": "矿种判断",
  "grade_info": "品位信息（如CaF2、SiO2）",
  "deposit_type": "成矿类型",
  "orebody_scale": "矿体规模",
  "thickness_extension": "厚度与延伸",
  "mineability": "可采性评估",
  "geological_info": "地质信息识别综合结论",
  "orebody_analysis": "矿体信息分析综合结论",
  "data_integrity": "数据完整性（是否具备决策所需信息、是否缺关键数据）",
  "risk_identification": "风险识别（地质风险/开采风险/合规风险）",
  "investment_advice": "建议继续/谨慎/放弃（三选一）",
  "result_interpretation": "结果说明",
  "logic_basis": "逻辑依据",
  "risk_hint": "风险提示",
  "risk_level": "低风险/中风险/高风险",
  "radar_scores": {{
    "geological_potential": 0-100,
    "data_integrity": 0-100,
    "project_stage": 0-100,
    "risk_factor": 0-100
  }},
  "highlight_keywords": ["矿种", "CaF2", "SiO2"]
}}
""".strip()

    return f"""
你是企业文档分析顾问。请输出严格JSON，不要输出JSON外文本。

{hard_rules}

公司名称：{company_name or '未提供'}
项目名称：{project_name or '未提供'}
文档内容：
{clipped}

输出JSON：
{{
  "summary": "文件摘要",
  "key_insights": ["关键信息1", "关键信息2", "关键信息3"],
  "bullet_points": ["结构化要点1", "结构化要点2", "结构化要点3"],
  "risk_issues": "风险/问题识别（若无则写未发现显著风险）",
  "application_advice": "应用建议（是否值得参考）",
  "result_interpretation": "结果说明",
  "logic_basis": "逻辑依据",
  "risk_hint": "风险提示",
  "risk_level": "低风险/中风险/高风险",
  "highlight_keywords": ["关键词1", "关键词2", "关键词3"]
}}
""".strip()


def _extract_json(text: str) -> dict:
    text = text.strip()
    if text.startswith("{") and text.endswith("}"):
        return json.loads(text)
    match = re.search(r"\{[\s\S]*\}", text)
    if not match:
        raise ValueError("未识别到JSON结构")
    return json.loads(match.group(0))


def _to_list(value: object) -> List[str]:
    if isinstance(value, list):
        return [str(v) for v in value]
    if isinstance(value, str) and value.strip():
        return [value.strip()]
    return []


def _safe_scores(payload: dict) -> Dict[str, float]:
    radar = payload.get("radar_scores", {}) or {}

    def _num(v: object, default: float) -> float:
        try:
            return float(v)
        except Exception:
            return default

    return {
        "geological_potential": max(0, min(100, _num(radar.get("geological_potential"), 60))),
        "data_integrity": max(0, min(100, _num(radar.get("data_integrity"), 55))),
        "project_stage": max(0, min(100, _num(radar.get("project_stage"), 50))),
        "risk_factor": max(0, min(100, _num(radar.get("risk_factor"), 55))),
    }


def analyze_file_with_ai(report_text: str, company_name: str = "", project_name: str = "") -> FileAIResult:
    if not report_text.strip():
        raise ValueError("文本为空，无法分析")

    mode = "mining" if _is_mining_related(report_text) else "general"
    raw = call_claude(_build_prompt(mode, company_name, project_name, report_text))

    try:
        payload = _extract_json(raw)
    except Exception as exc:
        raise RuntimeError("AI返回格式异常，请重试。") from exc

    if mode == "mining":
        advice_raw = str(payload.get("investment_advice", "谨慎"))
        if "继续" in advice_raw or "建议投资" in advice_raw:
            advice = "建议继续"
        elif "放弃" in advice_raw:
            advice = "放弃"
        else:
            advice = "谨慎"

        return FileAIResult(
            mode="mining",
            summary=str(payload.get("summary", "")),
            key_insights=[],
            bullet_points=[],
            risk_issues="",
            application_advice="",
            mineral_type=str(payload.get("mineral_type", "")),
            grade_info=str(payload.get("grade_info", "")),
            deposit_type=str(payload.get("deposit_type", "")),
            orebody_scale=str(payload.get("orebody_scale", "")),
            thickness_extension=str(payload.get("thickness_extension", "")),
            mineability=str(payload.get("mineability", "")),
            geological_info=str(payload.get("geological_info", "")),
            orebody_analysis=str(payload.get("orebody_analysis", "")),
            data_integrity=str(payload.get("data_integrity", "")),
            risk_identification=str(payload.get("risk_identification", "")),
            investment_advice=advice,
            result_interpretation=str(payload.get("result_interpretation", "")),
            logic_basis=str(payload.get("logic_basis", "")),
            risk_hint=str(payload.get("risk_hint", "")),
            risk_level=str(payload.get("risk_level", "中风险")),
            radar_scores=_safe_scores(payload),
            highlight_keywords=_to_list(payload.get("highlight_keywords")),
            raw_response=raw,
        )

    return FileAIResult(
        mode="general",
        summary=str(payload.get("summary", "")),
        key_insights=_to_list(payload.get("key_insights")),
        bullet_points=_to_list(payload.get("bullet_points")),
        risk_issues=str(payload.get("risk_issues", "")),
        application_advice=str(payload.get("application_advice", "")),
        mineral_type="",
        grade_info="",
        deposit_type="",
        orebody_scale="",
        thickness_extension="",
        mineability="",
        geological_info="",
        orebody_analysis="",
        data_integrity="",
        risk_identification="",
        investment_advice="",
        result_interpretation=str(payload.get("result_interpretation", "")),
        logic_basis=str(payload.get("logic_basis", "")),
        risk_hint=str(payload.get("risk_hint", "")),
        risk_level=str(payload.get("risk_level", "低风险")),
        radar_scores={"geological_potential": 0, "data_integrity": 0, "project_stage": 0, "risk_factor": 0},
        highlight_keywords=_to_list(payload.get("highlight_keywords")),
        raw_response=raw,
    )
